from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


def _wrap_text(text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
    """Wrap text for PDF by measuring string width.

    Works reasonably for both Chinese and English (char-level wrapping for Chinese).
    """
    # Normalize
    text = text.rstrip("\n")
    if not text:
        return [""]

    lines: list[str] = []
    current = ""
    for ch in text:
        if ch == "\n":
            lines.append(current)
            current = ""
            continue
        candidate = current + ch
        w = pdfmetrics.stringWidth(candidate, font_name, font_size)
        if w <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current or not lines:
        lines.append(current)
    return lines


def md_to_docx(md_text: str, out_path: Path, title: str) -> None:
    doc = Document()

    # Title
    doc.add_heading(title, level=0)

    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        code_buf = []

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
            else:
                in_code = False
                flush_code()
            continue

        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        m = re.match(r"^(#+)\s+(.*)$", line)
        if m:
            level = min(4, len(m.group(1)))
            text = m.group(2).strip()
            doc.add_heading(text, level=level)
            continue

        if line.lstrip().startswith(("- ", "* ")):
            text = line.lstrip()[2:].strip()
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Normal paragraph
        doc.add_paragraph(line)

    flush_code()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def md_to_pdf(md_text: str, out_path: Path, title: str, font_path: str) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    font_name = "CNFont"
    pdfmetrics.registerFont(TTFont(font_name, font_path))

    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4

    margin_x = 50
    margin_y = 50
    cursor_y = height - margin_y

    def new_page():
        nonlocal cursor_y
        c.showPage()
        c.setFont(font_name, 11)
        cursor_y = height - margin_y

    # Title
    c.setFont(font_name, 16)
    for l in _wrap_text(title, font_name, 16, width - 2 * margin_x):
        c.drawString(margin_x, cursor_y, l)
        cursor_y -= 22
    cursor_y -= 6
    c.setFont(font_name, 11)

    in_code = False
    code_font = "Courier"
    c.setFont(font_name, 11)

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if not line.strip():
            cursor_y -= 12
            if cursor_y < margin_y:
                new_page()
            continue

        # Headings
        m = re.match(r"^(#+)\s+(.*)$", line)
        if m and not in_code:
            level = len(m.group(1))
            text = m.group(2).strip()
            size = 14 if level == 1 else 12
            c.setFont(font_name, size)
            for l in _wrap_text(text, font_name, size, width - 2 * margin_x):
                c.drawString(margin_x, cursor_y, l)
                cursor_y -= size + 6
            cursor_y -= 2
            c.setFont(font_name, 11)
            if cursor_y < margin_y:
                new_page()
            continue

        # Bullets
        if line.lstrip().startswith(("- ", "* ")) and not in_code:
            text = line.lstrip()[2:].strip()
            bullet_prefix = "• "
            wrapped = _wrap_text(bullet_prefix + text, font_name, 11, width - 2 * margin_x)
            for i, l in enumerate(wrapped):
                c.drawString(margin_x, cursor_y, l)
                cursor_y -= 14
                if cursor_y < margin_y:
                    new_page()
            continue

        # Code / normal
        if in_code:
            c.setFont(code_font, 9)
            wrapped = _wrap_text(line, code_font, 9, width - 2 * margin_x)
            for l in wrapped:
                c.drawString(margin_x, cursor_y, l)
                cursor_y -= 12
                if cursor_y < margin_y:
                    new_page()
            c.setFont(font_name, 11)
        else:
            wrapped = _wrap_text(line, font_name, 11, width - 2 * margin_x)
            for l in wrapped:
                c.drawString(margin_x, cursor_y, l)
                cursor_y -= 14
                if cursor_y < margin_y:
                    new_page()

    c.save()


def main() -> None:
    root = Path(__file__).resolve().parent
    docs_dir = root / "docs"
    out_dir = root / "docs_export"
    out_dir.mkdir(parents=True, exist_ok=True)

    # A Chinese-capable font
    font_path = "/usr/share/fonts/truetype/arphic/uming.ttc"
    if not Path(font_path).exists():
        raise SystemExit(f"Chinese font not found: {font_path}")

    items = [
        (docs_dir / "PRD.md", "产品需求文档（PRD）- SwarmAI Platform", "产品需求文档_PRD"),
        (docs_dir / "Product_Design.md", "产品设计说明书（PDD）- SwarmAI Platform", "产品设计说明书_PDD"),
        (docs_dir / "API_Spec.md", "产品接口报告（API）- SwarmAI Platform", "产品接口报告_API"),
        (docs_dir / "Protocol_Spec.md", "产品接口报告（协议）- SwarmAI Platform", "产品接口报告_协议"),
        (docs_dir / "Deployment_Guide.md", "部署与运行说明 - SwarmAI Platform", "部署与运行说明"),
        (docs_dir / "User_Manual.md", "产品说明书（用户手册）- SwarmAI Platform", "产品说明书_用户手册"),
    ]

    for md_path, title, stem in items:
        if not md_path.exists():
            continue
        md_text = md_path.read_text(encoding="utf-8")
        md_to_docx(md_text, out_dir / f"{stem}.docx", title)
        md_to_pdf(md_text, out_dir / f"{stem}.pdf", title, font_path)

    print(f"Exported docs to: {out_dir}")


if __name__ == "__main__":
    main()
